# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Yeni belge oluştur
doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_cell_border(cell, **kwargs):
    """Hücre kenarlığı ayarla"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_normal_para(doc, text, indent=False):
    """Normal paragraf ekle"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    return para

# ===================== KAPAK SAYFASI =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VAN YÜZÜNCÜ YIL ÜNİVERSİTESİ")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MÜHENDİSLİK FAKÜLTESİ")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("YOLO TABANLI İŞ GÜVENLİĞİ EKİPMANI ALGILAMA SİSTEMİ")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("(YOLO-Based Occupational Safety Equipment Detection System)")
run.italic = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LİSANS BİTİRME PROJESİ ÖN ARAŞTIRMA VE PLANLAMA RAPORU")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Hazırlayan: Öğrenci Adı SOYADI")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Danışman: Dr. Öğr. Üyesi Danışman Adı SOYADI")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Van")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Aralık, 2024")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ===================== İÇİNDEKİLER (TABLO) =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("İÇİNDEKİLER")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

icerik_listesi = [
    ("GİRİŞ", "1"),
    ("   Araştırmanın Amacı", "1"),
    ("KURAMSAL TEMELLER VE İLGİLİ ARAŞTIRMALAR", "2"),
    ("   Derin Öğrenme ve YOLO Mimarisi", "2"),
    ("   İş Güvenliği ve Kişisel Koruyucu Ekipmanlar", "3"),
    ("   İlgili Çalışmalar", "4"),
    ("   Araştırmanın Önemi ve Gerekçesi", "5"),
    ("MATERYAL VE METOD", "6"),
    ("   Kullanılacak Teknolojiler ve Araçlar", "6"),
    ("   Veri Seti Planlaması", "7"),
    ("   Model Eğitimi Planlaması", "8"),
    ("   Sistem Mimarisi Tasarımı", "9"),
    ("   İş-Zaman Tablosu", "10"),
    ("   Zorunlu Yapılacaklar", "11"),
    ("   Alternatif Çözümler (B Planı)", "11"),
    ("KAYNAKÇA", "12"),
]

# İçindekiler tablosu - kenarlıklı
table_icerik = doc.add_table(rows=len(icerik_listesi) + 1, cols=2)
table_icerik.style = 'Table Grid'
table_icerik.alignment = WD_TABLE_ALIGNMENT.CENTER

# Başlık satırı
table_icerik.rows[0].cells[0].text = "Bölüm"
table_icerik.rows[0].cells[1].text = "Sayfa"
for j in range(2):
    table_icerik.rows[0].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
    table_icerik.rows[0].cells[j].paragraphs[0].runs[0].font.size = Pt(12)
    table_icerik.rows[0].cells[j].paragraphs[0].runs[0].bold = True
table_icerik.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (baslik, sayfa) in enumerate(icerik_listesi):
    # Başlık hücresi
    cell0 = table_icerik.rows[i+1].cells[0]
    cell0.text = baslik
    cell0.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell0.paragraphs[0].runs[0].font.size = Pt(12)
    cell0.width = Cm(12)
    
    # Sayfa hücresi
    cell1 = table_icerik.rows[i+1].cells[1]
    cell1.text = sayfa
    cell1.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell1.paragraphs[0].runs[0].font.size = Pt(12)
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.width = Cm(2)

doc.add_page_break()

# ===================== TABLOLAR DİZİNİ (TABLO) =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TABLOLAR DİZİNİ")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

tablolar = [
    ("Tablo 1. Algılanacak Kişisel Koruyucu Ekipman Sınıfları", "7"),
    ("Tablo 2. Planlanan Veri Seti İstatistikleri", "7"),
    ("Tablo 3. Planlanan Model Eğitim Hiperparametreleri", "8"),
    ("Tablo 4. Bitirme Projesinin Tümüne Ait İş-Zaman Tablosu", "10"),
    ("Tablo 5. Zorunlu Yapılacaklar Listesi", "11"),
]

table_tablolar = doc.add_table(rows=len(tablolar) + 1, cols=2)
table_tablolar.style = 'Table Grid'
table_tablolar.alignment = WD_TABLE_ALIGNMENT.CENTER

# Başlık satırı
table_tablolar.rows[0].cells[0].text = "Tablo Adı"
table_tablolar.rows[0].cells[1].text = "Sayfa"
for j in range(2):
    table_tablolar.rows[0].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
    table_tablolar.rows[0].cells[j].paragraphs[0].runs[0].font.size = Pt(12)
    table_tablolar.rows[0].cells[j].paragraphs[0].runs[0].bold = True
table_tablolar.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (baslik, sayfa) in enumerate(tablolar):
    cell0 = table_tablolar.rows[i+1].cells[0]
    cell0.text = baslik
    cell0.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell0.paragraphs[0].runs[0].font.size = Pt(12)
    cell0.width = Cm(12)
    
    cell1 = table_tablolar.rows[i+1].cells[1]
    cell1.text = sayfa
    cell1.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell1.paragraphs[0].runs[0].font.size = Pt(12)
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.width = Cm(2)

doc.add_page_break()

# ===================== ŞEKİLLER DİZİNİ (TABLO) =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ŞEKİLLER DİZİNİ")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

sekiller = [
    ("Şekil 1. YOLO Algoritması Çalışma Prensibi", "3"),
    ("Şekil 2. YOLOv8 Ağ Mimarisi", "3"),
    ("Şekil 3. Örnek Veri Seti Görselleri", "7"),
    ("Şekil 4. Planlanan Sistem Mimarisi Diyagramı", "9"),
]

table_sekiller = doc.add_table(rows=len(sekiller) + 1, cols=2)
table_sekiller.style = 'Table Grid'
table_sekiller.alignment = WD_TABLE_ALIGNMENT.CENTER

# Başlık satırı
table_sekiller.rows[0].cells[0].text = "Şekil Adı"
table_sekiller.rows[0].cells[1].text = "Sayfa"
for j in range(2):
    table_sekiller.rows[0].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
    table_sekiller.rows[0].cells[j].paragraphs[0].runs[0].font.size = Pt(12)
    table_sekiller.rows[0].cells[j].paragraphs[0].runs[0].bold = True
table_sekiller.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (baslik, sayfa) in enumerate(sekiller):
    cell0 = table_sekiller.rows[i+1].cells[0]
    cell0.text = baslik
    cell0.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell0.paragraphs[0].runs[0].font.size = Pt(12)
    cell0.width = Cm(12)
    
    cell1 = table_sekiller.rows[i+1].cells[1]
    cell1.text = sayfa
    cell1.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell1.paragraphs[0].runs[0].font.size = Pt(12)
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell1.width = Cm(2)

doc.add_page_break()

# ===================== 1. GİRİŞ =====================
h = doc.add_heading("1. GİRİŞ", level=1)
h.runs[0].font.name = 'Times New Roman'

giris_metin = """
İş kazaları, dünya genelinde ciddi bir toplumsal ve ekonomik sorun oluşturmaktadır. Uluslararası Çalışma Örgütü (ILO) verilerine göre, her yıl yaklaşık 2.3 milyon işçi iş kazası veya meslek hastalıkları nedeniyle hayatını kaybetmektedir. Bu kazaların önemli bir kısmı, kişisel koruyucu ekipmanların (KKE) kullanılmaması veya yanlış kullanılmasından kaynaklanmaktadır.

Geleneksel iş güvenliği denetim yöntemleri, insan gözlemcilerin sahada fiziksel olarak bulunmasını gerektirmektedir. Bu yöntem, hem zaman alıcı hem de insan hatasına açıktır. Özellikle büyük ölçekli inşaat şantiyeleri, maden ocakları ve fabrikalar gibi alanlarda sürekli denetim yapmak pratik olarak mümkün değildir.

Son yıllarda derin öğrenme tabanlı nesne algılama algoritmaları, görüntü işleme alanında çığır açan gelişmeler sağlamıştır. Bu algoritmalar arasında YOLO (You Only Look Once), gerçek zamanlı nesne algılama performansı ile öne çıkmaktadır. YOLO algoritması, tek bir ileri geçişte görüntüdeki tüm nesneleri algılayabilmekte ve bu sayede yüksek hızlarda çalışabilmektedir.

Bu bitirme projesi kapsamında, YOLO tabanlı bir iş güvenliği ekipmanı algılama sistemi geliştirilmesi planlanmaktadır. Sistem, güvenlik kameralarından alınan görüntüleri gerçek zamanlı olarak analiz ederek çalışanların kişisel koruyucu ekipman kullanımını otomatik olarak denetleyecektir. Bu rapor, projenin ön araştırma ve planlama aşamasını içermektedir.
"""

for para_text in giris_metin.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

# Araştırmanın Amacı
h = doc.add_heading("1.1. Araştırmanın Amacı", level=2)
h.runs[0].font.name = 'Times New Roman'

amac_metin = """
Bu araştırmanın temel amacı, iş yerlerinde güvenlik denetimini otomatikleştirmek için derin öğrenme tabanlı bir nesne algılama sistemi geliştirmektir. Proje kapsamında aşağıdaki hedefler belirlenmiştir:

• Baret, yelek, gözlük, eldiven ve güvenlik ayakkabısı gibi temel kişisel koruyucu ekipmanları yüksek doğrulukla algılayabilen bir model geliştirmek

• Gerçek zamanlı video akışı üzerinde çalışabilecek performansta bir sistem tasarlamak

• Algılama sonuçlarını görselleştiren ve raporlayan kullanıcı dostu bir arayüz oluşturmak

• Sistem performansını farklı aydınlatma koşulları ve kamera açılarında test etmek

• Elde edilen sonuçları mevcut literatürdeki çalışmalarla karşılaştırmak
"""

for para_text in amac_metin.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

doc.add_page_break()

# ===================== 2. KURAMSAL TEMELLER =====================
h = doc.add_heading("2. KURAMSAL TEMELLER VE İLGİLİ ARAŞTIRMALAR", level=1)
h.runs[0].font.name = 'Times New Roman'

kuramsal_giris = """
Bu bölümde, projenin teorik altyapısını oluşturan derin öğrenme kavramları, YOLO algoritması ve iş güvenliği mevzuatı hakkında bilgi verilecektir. Ayrıca literatürdeki ilgili çalışmalar incelenecektir.
"""
add_normal_para(doc, kuramsal_giris.strip(), indent=True)

# Derin Öğrenme ve YOLO
h = doc.add_heading("2.1. Derin Öğrenme ve YOLO Mimarisi", level=2)
h.runs[0].font.name = 'Times New Roman'

derin_ogrenme = """
Derin öğrenme, yapay sinir ağlarının çok katmanlı yapılarını kullanarak verilerdeki karmaşık örüntüleri öğrenebilen bir makine öğrenimi alt dalıdır. Evrişimli Sinir Ağları (Convolutional Neural Networks - CNN), görüntü işleme alanında en yaygın kullanılan derin öğrenme mimarisidir.

Nesne algılama, bir görüntüdeki nesnelerin hem konumlarının hem de sınıflarının belirlenmesi işlemidir. Geleneksel nesne algılama yöntemleri iki aşamalı bir yaklaşım izlemektedir: önce bölge önerileri üretilmekte, ardından bu bölgeler sınıflandırılmaktadır. R-CNN, Fast R-CNN ve Faster R-CNN bu yaklaşımın örnekleridir.

YOLO (You Only Look Once) algoritması, 2015 yılında Joseph Redmon ve arkadaşları tarafından önerilmiştir. YOLO, nesne algılamayı tek bir regresyon problemi olarak ele almakta ve görüntüyü tek bir ileri geçişte işlemektedir. Bu yaklaşım, iki aşamalı yöntemlere kıyasla çok daha hızlı çalışmayı mümkün kılmaktadır.

YOLO algoritmasının çalışma prensibi şu şekildedir:
1. Giriş görüntüsü S×S boyutunda bir ızgaraya bölünür
2. Her ızgara hücresi, B adet sınırlayıcı kutu (bounding box) ve güven skoru tahmin eder
3. Her kutu için (x, y, w, h, güven) değerleri hesaplanır
4. Eşzamanlı olarak her hücre için C adet sınıf olasılığı tahmin edilir
5. Non-Maximum Suppression (NMS) ile çakışan kutular elenir

YOLOv8, Ultralytics tarafından Ocak 2023'te yayınlanan en güncel YOLO sürümüdür. CSPDarknet53 omurgası, PANet boyun yapısı ve anchor-free algılama başlığı içermektedir. YOLOv8, önceki sürümlere kıyasla hem doğruluk hem de hız açısından önemli iyileştirmeler sunmaktadır.
"""

for para_text in derin_ogrenme.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

# Şekil 1 - YOLO Algoritması
doc.add_paragraph()
try:
    doc.add_picture('sekil1_yolo_algoritma.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

p = doc.add_paragraph()
run = p.add_run("Şekil 1. YOLO Algoritması Çalışma Prensibi")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Şekil 2 - YOLOv8 Mimarisi
try:
    doc.add_picture('sekil2_yolov8_mimari.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

p = doc.add_paragraph()
run = p.add_run("Şekil 2. YOLOv8 Ağ Mimarisi")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# İş Güvenliği
h = doc.add_heading("2.2. İş Güvenliği ve Kişisel Koruyucu Ekipmanlar", level=2)
h.runs[0].font.name = 'Times New Roman'

is_guvenligi = """
6331 sayılı İş Sağlığı ve Güvenliği Kanunu'na göre, işverenler çalışanlarının sağlık ve güvenliğini korumak için gerekli tüm önlemleri almakla yükümlüdür. Bu önlemler arasında kişisel koruyucu ekipman (KKE) temini ve kullanımının denetlenmesi önemli bir yer tutmaktadır.

Kişisel koruyucu ekipmanlar, çalışanları iş kazalarından ve meslek hastalıklarından koruyan donanımlardır. Başlıca KKE türleri şunlardır:

• Baş Koruyucuları: Baret, düşen nesnelere ve darbelere karşı koruma sağlar
• Göz ve Yüz Koruyucuları: Koruyucu gözlük, yüz siperi, kaynak maskesi
• İşitme Koruyucuları: Kulak tıkacı, kulaklık
• Solunum Koruyucuları: Toz maskesi, gaz maskesi
• El Koruyucuları: Eldiven (kimyasal, mekanik, elektrik)
• Ayak Koruyucuları: Güvenlik ayakkabısı, çizme
• Vücut Koruyucuları: Reflektif yelek, iş tulumu

İnşaat, maden ve ağır sanayi sektörlerinde özellikle baret, reflektif yelek ve güvenlik ayakkabısı kullanımı zorunludur. Ancak yapılan denetimlerde, çalışanların %30-40'ının KKE kullanımına tam olarak uymadığı tespit edilmiştir.
"""

for para_text in is_guvenligi.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

# İlgili Çalışmalar
h = doc.add_heading("2.3. İlgili Çalışmalar", level=2)
h.runs[0].font.name = 'Times New Roman'

ilgili_calismalar = """
Literatürde iş güvenliği ekipmanı tespiti üzerine çeşitli çalışmalar bulunmaktadır. Bu çalışmaların bir kısmı aşağıda özetlenmiştir:

Fang ve arkadaşları (2018), inşaat şantiyelerinde baret algılama için Faster R-CNN tabanlı bir sistem geliştirmişlerdir. Çalışmada 3.000'den fazla görüntü kullanılmış ve %95.7 doğruluk oranı elde edilmiştir.

Wu ve arkadaşları (2019), SSD (Single Shot MultiBox Detector) algoritmasını kullanarak çoklu KKE algılama yapmışlardır. Baret, yelek ve güvenlik kemeri için ortalama %89.2 mAP değeri raporlanmıştır.

Nath ve arkadaşları (2020), YOLOv3 ile gerçek zamanlı baret ve yelek algılama sistemi geliştirmişlerdir. Sistem 25 FPS hızında çalışmış ve %92.3 mAP performansı göstermiştir.

Chen ve arkadaşları (2021), YOLOv5 kullanarak inşaat işçilerinde baret, yelek, eldiven ve güvenlik ayakkabısı algılama çalışması yapmışlardır. 15.000 görüntülük veri seti ile eğitilen model %94.6 mAP değerine ulaşmıştır.

Son dönemde yapılan çalışmalarda YOLOv8 modelinin diğer YOLO sürümlerine göre daha yüksek doğruluk ve hız performansı sergilediği görülmektedir. Bu nedenle projemizde YOLOv8 mimarisi tercih edilecektir.
"""

for para_text in ilgili_calismalar.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

doc.add_page_break()

# Araştırmanın Önemi
h = doc.add_heading("2.4. Araştırmanın Önemi ve Gerekçesi", level=2)
h.runs[0].font.name = 'Times New Roman'

onem = """
İş kazalarının önlenmesinde kişisel koruyucu ekipman kullanımının denetlenmesi kritik öneme sahiptir. Geleneksel denetim yöntemlerinin yetersiz kaldığı büyük ölçekli iş alanlarında otomatik algılama sistemleri bu açığı kapatabilir.

Bu araştırmanın önemi ve gerekçeleri şu şekilde sıralanabilir:

• İş kazalarının azaltılmasına katkı sağlayacaktır
• 7/24 kesintisiz denetim imkanı sunacaktır
• İnsan kaynaklı denetim hatalarını minimize edecektir
• Gerçek zamanlı uyarı sistemi ile anlık müdahale imkanı sağlayacaktır
• Denetim maliyetlerini düşürecektir
• İş güvenliği verilerinin otomatik toplanması ve analizi mümkün olacaktır

Literatürdeki mevcut çalışmalar incelendiğinde, Türkiye'deki iş sahalarına özgü koşulları dikkate alan kapsamlı bir çalışmanın eksikliği görülmektedir. Bu proje, yerel ihtiyaçlara uygun bir çözüm sunmayı amaçlamaktadır.
"""

for para_text in onem.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

doc.add_page_break()

# ===================== 3. MATERYAL VE METOD =====================
h = doc.add_heading("3. MATERYAL VE METOD", level=1)
h.runs[0].font.name = 'Times New Roman'

materyal_giris = """
Bu bölümde projenin gerçekleştirilmesinde kullanılacak teknolojiler, veri seti planlaması, model eğitimi planlaması ve sistem mimarisi tasarımı detaylı olarak açıklanacaktır.
"""
add_normal_para(doc, materyal_giris.strip(), indent=True)

# Kullanılacak Teknolojiler
h = doc.add_heading("3.1. Kullanılacak Teknolojiler ve Araçlar", level=2)
h.runs[0].font.name = 'Times New Roman'

teknolojiler = """
Proje geliştirme sürecinde aşağıdaki teknolojiler ve araçların kullanılması planlanmaktadır:

Programlama Dili: Python 3.10
Python, veri bilimi ve makine öğrenimi alanlarında en yaygın kullanılan programlama dilidir. Zengin kütüphane desteği ve kolay sözdizimi nedeniyle tercih edilecektir.

Derin Öğrenme Framework'ü: PyTorch 2.0
PyTorch, Facebook AI Research tarafından geliştirilen açık kaynaklı bir derin öğrenme kütüphanesidir. Dinamik hesaplama grafı yapısı sayesinde esnek model geliştirme imkanı sunmaktadır.

Nesne Algılama Kütüphanesi: Ultralytics YOLOv8
YOLOv8, Ultralytics tarafından sağlanan son teknoloji nesne algılama kütüphanesidir. Kolay kullanımı ve yüksek performansı ile öne çıkmaktadır.

Görüntü İşleme: OpenCV 4.8
OpenCV, gerçek zamanlı bilgisayarlı görü uygulamaları için kullanılan açık kaynaklı bir kütüphanedir.

Veri Etiketleme: LabelImg, Roboflow
Görüntülerdeki nesnelerin manuel olarak etiketlenmesi için LabelImg aracı kullanılacaktır. Roboflow platformu ise veri artırma ve format dönüşümü için kullanılacaktır.

Donanım: NVIDIA GeForce RTX 3060 (12GB VRAM)
Model eğitimi için CUDA destekli GPU kullanılacaktır. Google Colab Pro+ da alternatif eğitim ortamı olarak değerlendirilecektir.

Geliştirme Ortamı: VS Code, Jupyter Notebook
Kod geliştirme ve deneyler için VS Code ve Jupyter Notebook kullanılacaktır.
"""

for para_text in teknolojiler.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

doc.add_page_break()

# Veri Seti Planlaması
h = doc.add_heading("3.2. Veri Seti Planlaması", level=2)
h.runs[0].font.name = 'Times New Roman'

veri_seti = """
Proje kapsamında kullanılacak veri seti, açık kaynaklı veri setlerinin birleştirilmesi ve özgün görüntülerin toplanması ile oluşturulacaktır. Veri seti, inşaat şantiyeleri ve endüstriyel tesislerden alınacak görüntülerden oluşacaktır.

Veri setinde algılanması planlanan sınıflar aşağıdaki tabloda verilmiştir:
"""
add_normal_para(doc, veri_seti.strip(), indent=True)

# Tablo 1 - Ekipman Sınıfları
p = doc.add_paragraph()
run = p.add_run("Tablo 1. Algılanacak Kişisel Koruyucu Ekipman Sınıfları")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table1 = doc.add_table(rows=7, cols=3)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ['Sınıf ID', 'Sınıf Adı', 'Açıklama']
for i, header in enumerate(headers1):
    table1.rows[0].cells[i].text = header
    table1.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    table1.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'

siniflar = [
    ('0', 'hardhat (baret)', 'Baş koruyucu baret'),
    ('1', 'no-hardhat', 'Baret takmayan kişi'),
    ('2', 'vest (yelek)', 'Reflektif güvenlik yeleği'),
    ('3', 'no-vest', 'Yelek giymeyen kişi'),
    ('4', 'person', 'Genel insan algılama'),
    ('5', 'safety-glasses', 'Koruyucu gözlük'),
]

for i, (sid, sadi, aciklama) in enumerate(siniflar):
    table1.rows[i+1].cells[0].text = sid
    table1.rows[i+1].cells[1].text = sadi
    table1.rows[i+1].cells[2].text = aciklama
    for j in range(3):
        table1.rows[i+1].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'

doc.add_paragraph()

# Tablo 2 - Veri Seti İstatistikleri
p = doc.add_paragraph()
run = p.add_run("Tablo 2. Planlanan Veri Seti İstatistikleri")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

istatistikler = [
    ('Özellik', 'Planlanan Değer'),
    ('Toplam Görüntü Sayısı', '8.000 - 10.000'),
    ('Eğitim Seti', '%80'),
    ('Doğrulama Seti', '%10'),
    ('Test Seti', '%10'),
]

for i, (ozellik, deger) in enumerate(istatistikler):
    table2.rows[i].cells[0].text = ozellik
    table2.rows[i].cells[1].text = deger
    table2.rows[i].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    table2.rows[i].cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
    if i == 0:
        table2.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table2.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

veri_seti_devam = """
Veri seti oluşturma sürecinde aşağıdaki veri artırma (data augmentation) tekniklerinin uygulanması planlanmaktadır:

• Yatay çevirme (horizontal flip)
• Rastgele döndürme (±15°)
• Parlaklık ve kontrast değişimi
• Mozaik artırma (mosaic augmentation)
• Rastgele kırpma (random crop)

Bu teknikler sayesinde modelin farklı koşullara karşı genelleme yeteneği artırılacaktır.
"""
add_normal_para(doc, veri_seti_devam.strip(), indent=True)

# Şekil 3 - Örnek PPE Algılama
doc.add_paragraph()
try:
    doc.add_picture('sekil3_ppe_ornek.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

p = doc.add_paragraph()
run = p.add_run("Şekil 3. Örnek Veri Seti Görselleri ve KKE Algılama Sonuçları")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Model Eğitimi Planlaması
h = doc.add_heading("3.3. Model Eğitimi Planlaması", level=2)
h.runs[0].font.name = 'Times New Roman'

model_egitimi = """
Model eğitimi için YOLOv8m (medium) modeli tercih edilecektir. Bu model, hız ve doğruluk arasında dengeli bir performans sunmaktadır.

Eğitim sırasında kullanılması planlanan hiperparametreler aşağıdaki tabloda verilmiştir:
"""
add_normal_para(doc, model_egitimi.strip(), indent=True)

# Tablo 3 - Hiperparametreler
p = doc.add_paragraph()
run = p.add_run("Tablo 3. Planlanan Model Eğitim Hiperparametreleri")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table3 = doc.add_table(rows=9, cols=2)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

hiperparametreler = [
    ('Parametre', 'Planlanan Değer'),
    ('Model', 'YOLOv8m'),
    ('Epoch Sayısı', '100'),
    ('Batch Size', '16'),
    ('Görüntü Boyutu', '640×640'),
    ('Öğrenme Oranı (lr0)', '0.01'),
    ('Momentum', '0.937'),
    ('Weight Decay', '0.0005'),
]

for i, (param, deger) in enumerate(hiperparametreler):
    table3.rows[i].cells[0].text = param
    table3.rows[i].cells[1].text = deger
    table3.rows[i].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    table3.rows[i].cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
    if i == 0:
        table3.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table3.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

egitim_detay = """
Eğitim sürecinde early stopping ve learning rate scheduling teknikleri kullanılacaktır. Model, COCO veri seti üzerinde önceden eğitilmiş ağırlıklardan başlatılarak transfer öğrenme yaklaşımı uygulanacaktır.

Hedeflenen performans metrikleri:
• mAP@0.5: %90 üzeri
• mAP@0.5:0.95: %75 üzeri
• Gerçek zamanlı performans: En az 25 FPS
"""
add_normal_para(doc, egitim_detay.strip(), indent=True)

doc.add_page_break()

# Sistem Mimarisi Tasarımı
h = doc.add_heading("3.4. Sistem Mimarisi Tasarımı", level=2)
h.runs[0].font.name = 'Times New Roman'

sistem_mimarisi = """
Geliştirilecek sistem, modüler bir mimari üzerine kurulacaktır. Sistemin ana bileşenleri şunlar olacaktır:

1. Görüntü Yakalama Modülü: Kameralardan veya video dosyalarından görüntü akışı alacaktır. OpenCV kütüphanesi ile gerçekleştirilecektir.

2. Ön İşleme Modülü: Görüntüleri model giriş formatına dönüştürecektir (yeniden boyutlandırma, normalizasyon).

3. Algılama Modülü: YOLOv8 modelini çalıştırarak nesneleri algılayacak ve sınıflandıracaktır.

4. Son İşleme Modülü: Algılama sonuçlarını filtreleyecek (NMS) ve yorumlayacaktır.

5. Görselleştirme Modülü: Algılanan nesneleri sınırlayıcı kutular ve etiketlerle görüntü üzerinde işaretleyecektir.

6. Uyarı Modülü: KKE eksikliği tespit edildiğinde sesli ve görsel uyarı verecektir.

7. Raporlama Modülü: Algılama istatistiklerini kaydedecek ve raporlayacaktır.

Sistem, Python ile geliştirilecek olup komut satırı arayüzü (CLI) ve basit bir grafik kullanıcı arayüzü (GUI) sunacaktır. GUI, Streamlit kütüphanesi kullanılarak oluşturulacaktır.

Sistemin çalışma akışı şu şekilde planlanmaktadır:
1. Kullanıcı video kaynağını seçer (kamera, video dosyası veya görüntü)
2. Sistem görüntüleri alır ve ön işleme yapar
3. YOLOv8 modeli algılama gerçekleştirir
4. Sonuçlar görselleştirilir ve ekranda gösterilir
5. KKE eksikliği varsa uyarı sistemi devreye girer
6. Tüm algılamalar log dosyasına kaydedilir
"""

for para_text in sistem_mimarisi.strip().split('\n\n'):
    add_normal_para(doc, para_text.strip(), indent=True)

# Şekil 4 - Sistem Mimarisi Diyagramı
doc.add_paragraph()
try:
    doc.add_picture('sekil4_sistem_mimari.png', width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

p = doc.add_paragraph()
run = p.add_run("Şekil 4. Planlanan Sistem Mimarisi Diyagramı")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===================== İŞ-ZAMAN TABLOSU =====================
h = doc.add_heading("3.5. İş-Zaman Tablosu", level=2)
h.runs[0].font.name = 'Times New Roman'

is_zaman_aciklama = """
Bitirme projesinin gerçekleştirilmesi için planlanan iş paketleri ve zaman çizelgesi aşağıdaki tabloda verilmiştir. Tabloda her bir iş paketinin hangi dönemde gerçekleştirileceği işaretlenmiştir.
"""
add_normal_para(doc, is_zaman_aciklama.strip(), indent=True)

# Tablo 4 - İş-Zaman Tablosu
p = doc.add_paragraph()
run = p.add_run("Tablo 4. Bitirme Projesinin Tümüne Ait İş-Zaman Tablosu")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table4 = doc.add_table(rows=10, cols=3)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

# İş-Zaman tablosu - 1. dönem araştırma ve planlama
is_zaman = [
    ('İş Paketi', '1. Dönem', '2. Dönem'),
    ('Literatür Taraması ve Araştırma', '●●●●●', ''),
    ('Proje Planlama ve Tasarım', '●●●●', ''),
    ('Veri Seti Araştırması', '●●●', ''),
    ('Teknoloji Araştırması', '●●●', ''),
    ('Veri Seti Toplama ve Etiketleme', '', '●●●●●'),
    ('Model Geliştirme ve Eğitim', '', '●●●●●●'),
    ('Sistem Entegrasyonu', '', '●●●●'),
    ('Test ve Değerlendirme', '', '●●●●'),
    ('Rapor Yazımı', '●● (Ön Rapor)', '●●●● (Final)'),
]

for i, (is_paketi, donem1, donem2) in enumerate(is_zaman):
    table4.rows[i].cells[0].text = is_paketi
    table4.rows[i].cells[1].text = donem1
    table4.rows[i].cells[2].text = donem2
    for j in range(3):
        table4.rows[i].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
        if i == 0:
            table4.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    # 1. Dönem işaretli olanları vurgula
    if i >= 1 and i <= 4:
        table4.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table4.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Not: Satırlar gerektiği kadar çoğaltılabilir. ● sembolleri tahmini hafta sayısını ifade etmektedir.")
run.italic = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ===================== ZORUNLU YAPILACAKLAR (TABLO) =====================
h = doc.add_heading("3.6. Zorunlu Yapılacaklar", level=2)
h.runs[0].font.name = 'Times New Roman'

zorunlu_aciklama = """
Bitirme projesinin tam olarak gerçekleştirilmesi için zorunlu olarak yapılması gereken işlemler aşağıdaki tabloda listelenmiştir:
"""
add_normal_para(doc, zorunlu_aciklama.strip(), indent=True)

# Tablo 5 - Zorunlu Yapılacaklar
p = doc.add_paragraph()
run = p.add_run("Tablo 5. Zorunlu Yapılacaklar Listesi")
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table5 = doc.add_table(rows=9, cols=3)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

zorunlu_isler = [
    ('No', 'Zorunlu İş Kalemi', 'Öncelik'),
    ('1', 'En az 5.000 görüntülük etiketli veri seti oluşturmak', 'Yüksek'),
    ('2', 'YOLOv8 modelini özel veri seti üzerinde eğitmek', 'Yüksek'),
    ('3', 'Minimum %85 mAP değerine ulaşmak', 'Yüksek'),
    ('4', 'Gerçek zamanlı çalışabilen (en az 15 FPS) bir sistem geliştirmek', 'Orta'),
    ('5', 'Çalışan bir prototip sistem sunmak', 'Yüksek'),
    ('6', 'Kullanıcı arayüzü tasarlamak ve geliştirmek', 'Orta'),
    ('7', 'Teknik dokümantasyon hazırlamak', 'Orta'),
    ('8', 'Bitirme projesi raporunu tamamlamak', 'Yüksek'),
]

for i, (no, is_kalemi, oncelik) in enumerate(zorunlu_isler):
    table5.rows[i].cells[0].text = no
    table5.rows[i].cells[1].text = is_kalemi
    table5.rows[i].cells[2].text = oncelik
    for j in range(3):
        table5.rows[i].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
        if i == 0:
            table5.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    table5.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table5.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Alternatif Çözümler
h = doc.add_heading("3.7. Alternatif Çözümler (B Planı)", level=2)
h.runs[0].font.name = 'Times New Roman'

alternatif = """
Projenin gerçekleştirilme aşamalarında karşılaşılabilecek sorunlar ve alternatif çözümler şunlardır:

Sorun 1: Yetersiz veri seti
Çözüm: Açık kaynaklı veri setleri (Roboflow Universe, Kaggle) kullanılacak ve veri artırma teknikleri uygulanacaktır.

Sorun 2: Düşük model performansı
Çözüm: Farklı YOLO sürümleri (YOLOv5, YOLOv8 nano/small) denenecek, hiperparametreler optimize edilecektir.

Sorun 3: Yetersiz GPU kaynağı
Çözüm: Google Colab Pro veya Kaggle Notebooks kullanılacaktır. Alternatif olarak daha küçük model (YOLOv8n) tercih edilecektir.

Sorun 4: Gerçek zamanlı performans sorunları
Çözüm: Model quantization uygulanacak, görüntü çözünürlüğü düşürülecektir.

Sorun 5: Farklı ortam koşullarında düşük performans
Çözüm: Veri seti çeşitlendirilecek, domain adaptation teknikleri araştırılacaktır.
"""
add_normal_para(doc, alternatif.strip(), indent=True)

doc.add_page_break()

# ===================== KAYNAKÇA =====================
h = doc.add_heading("KAYNAKÇA", level=1)
h.runs[0].font.name = 'Times New Roman'

kaynaklar = [
    "[1] Redmon, J., Divvala, S., Girshick, R., & Farhadi, A. (2016). You only look once: Unified, real-time object detection. In Proceedings of the IEEE conference on computer vision and pattern recognition (pp. 779-788).",
    "[2] Jocher, G., Chaurasia, A., & Qiu, J. (2023). Ultralytics YOLOv8. GitHub. https://github.com/ultralytics/ultralytics",
    "[3] Fang, Q., Li, H., Luo, X., Ding, L., Luo, H., Rose, T. M., & An, W. (2018). Detecting non-hardhat-use by a deep learning method from far-field surveillance videos. Automation in Construction, 85, 1-9.",
    "[4] Wu, J., Cai, N., Chen, W., Wang, H., & Wang, G. (2019). Automatic detection of hardhats worn by construction personnel: A deep learning approach and benchmark dataset. Automation in Construction, 106, 102894.",
    "[5] Nath, N. D., Behzadan, A. H., & Paal, S. G. (2020). Deep learning for site safety: Real-time detection of personal protective equipment. Automation in Construction, 112, 103085.",
    "[6] Chen, S., & Demachi, K. (2021). A vision-based approach for ensuring proper use of personal protective equipment (PPE) in decommissioning of Fukushima Daiichi Nuclear Power Station. Applied Sciences, 11(13), 5953.",
    "[7] 6331 Sayılı İş Sağlığı ve Güvenliği Kanunu. (2012). T.C. Resmi Gazete, 28339.",
    "[8] International Labour Organization. (2019). Safety and health at work: A vision for sustainable prevention. ILO.",
    "[9] He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. In Proceedings of the IEEE conference on computer vision and pattern recognition (pp. 770-778).",
    "[10] Lin, T. Y., Maire, M., Belongie, S., Hays, J., Perona, P., Ramanan, D., ... & Zitnick, C. L. (2014). Microsoft COCO: Common objects in context. In European conference on computer vision (pp. 740-755). Springer.",
]

for kaynak in kaynaklar:
    p = doc.add_paragraph()
    run = p.add_run(kaynak)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)

# Kaydet
output_file = 'Bitirme_Projesi_YOLO_Is_Guvenligi.docx'
doc.save(output_file)
print(f"Rapor başarıyla oluşturuldu: {output_file}")

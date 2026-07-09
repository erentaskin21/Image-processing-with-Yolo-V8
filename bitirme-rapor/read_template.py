from docx import Document

doc = Document('Bitirme_Projesi_rapor_Sablonu.docx')

print("=== PARAGRAFLAR ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text[:200]}")

print("\n=== TABLOLAR ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- Tablo {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        row_text = [cell.text[:50] for cell in row.cells]
        print(f"  Satır {r_idx}: {row_text}")
